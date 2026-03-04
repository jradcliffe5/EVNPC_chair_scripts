#!/usr/bin/env python3
"""Rename chair review files based on CSV records.

The script expects the CSV to provide `Name` and `Surname` columns (as in the
Google Forms export). For each row it looks in the source directory for a file
matching the pattern:

    <input_name> - Name Surname.txt

and moves it to the destination directory with the new filename:

    <prefix>_Name_Surname.txt

Both the original input name and the target prefix are case-sensitive strings,
while reviewer names are matched case-insensitively (with redundant whitespace
collapsed). The original file extension is preserved.

Example:
    python rename_reviews_from_csv.py pc_chair/EVN_PC_review_submission.csv \
        --prefix E25A001 --source-dir "Copy of EVN.../Review submission (File responses)" \
        --dest-dir pc_chair/reviews
    python rename_reviews_from_csv.py "https://docs.google.com/spreadsheets/d/.../edit#gid=0" \
        --prefix E25A001 --source-dir downloads --dest-dir pc_chair/reviews
"""

from __future__ import annotations

import argparse
import csv
import io
import re
import shutil
import sys
import unicodedata
import urllib.error
import urllib.request
from contextlib import contextmanager
from pathlib import Path
from textwrap import dedent
from typing import Iterable, Optional, Sequence, Tuple


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Rename existing review files from '<input> - Name Surname.ext' to "
            "'<prefix>_Name_Surname.ext'."
        ),
        epilog=dedent(
            """\
            Examples:
              rename_reviews_from_csv.py pc_chair/EVN_PC_review_submission.csv --prefix E25A001
              rename_reviews_from_csv.py responses.csv --prefix E25A002 --source-dir downloads --dest-dir reviews --dry-run
            """
        ),
    )
    parser.add_argument(
        "csv_path",
        help="Path to the CSV export or a Google Sheets URL.",
    )
    parser.add_argument(
        "--prefix",
        required=True,
        help="Text to prepend to each output filename (replaces <input> placeholder).",
    )
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=None,
        help=(
            "Directory containing the downloaded review files. "
            "Defaults to the CSV's parent directory (or the current working directory for URLs)."
        ),
    )
    parser.add_argument(
        "--dest-dir",
        type=Path,
        default=Path.cwd(),
        help="Directory where renamed files will be placed (created if missing).",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show the planned moves without changing any files.",
    )
    parser.add_argument(
        "--skip-missing",
        action="store_true",
        help="Silently continue when a referenced file cannot be located.",
    )
    parser.add_argument(
        "--no-docx-conversion",
        action="store_true",
        help="Disable automatic conversion of .docx submissions to .txt.",
    )
    parser.add_argument(
        "--prefer-newest",
        action="store_true",
        help=(
            "When multiple submissions match a reviewer, choose the most recently "
            "modified file."
        ),
    )
    return parser.parse_args(argv)


def sanitized_component(value: str) -> str:
    """Return a filesystem-safe name component derived from the input."""
    stripped = value.strip()
    if not stripped:
        return "Unknown"
    filtered = "".join(ch for ch in stripped if ch.isalnum())
    return filtered or "Unknown"


def normalize_name(value: str) -> str:
    """Normalise a name for comparison (case-insensitive, single spaces)."""
    normalized = unicodedata.normalize("NFC", value)
    return " ".join(normalized.lower().split())


def split_version_suffix(stem: str) -> Tuple[str, Optional[int]]:
    """Split common duplicate suffixes like ' (1)' or '_1' from a filename stem."""
    match = re.search(r"(?:\s*\((\d+)\)|_(\d+))$", stem)
    if not match:
        return stem, None
    version = int(match.group(1) or match.group(2))
    base = stem[: match.start()].rstrip()
    return base, version


def is_url(value: str) -> bool:
    return value.startswith(("http://", "https://"))


def gsheet_csv_url(url: str) -> str:
    """Return a CSV export URL for a Google Sheets share link."""
    if "format=csv" in url or "output=csv" in url:
        return url
    match = re.search(r"/d/([a-zA-Z0-9-_]+)", url)
    if not match:
        return url
    sheet_id = match.group(1)
    gid_match = re.search(r"[?&]gid=([0-9]+)", url)
    gid = gid_match.group(1) if gid_match else None
    export = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
    if gid:
        export = f"{export}&gid={gid}"
    return export


@contextmanager
def open_csv_source(csv_path: str) -> Iterable[Tuple[Iterable[str], str, Optional[Path]]]:
    """Yield a file-like object, label, and optional parent directory."""
    if is_url(csv_path):
        csv_url = gsheet_csv_url(csv_path)
        try:
            with urllib.request.urlopen(csv_url) as response:
                data = response.read()
        except urllib.error.URLError as exc:
            raise RuntimeError(f"Unable to download CSV from {csv_url}: {exc}") from exc
        text = data.decode("utf-8-sig", errors="replace")
        handle = io.StringIO(text)
        try:
            yield handle, csv_url, None
        finally:
            handle.close()
    else:
        path = Path(csv_path)
        if not path.is_file():
            raise FileNotFoundError(path)
        with path.open(newline="", encoding="utf-8") as handle:
            yield handle, str(path), path.parent


def find_submission_file(
    name: str,
    surname: str,
    source_dir: Path,
    prefer_newest: bool = False,
) -> Optional[Path]:
    """Locate a file matching '* - Name Surname.*' in the source directory."""
    target_tail = normalize_name(f"- {name} {surname}")
    matches: list[Tuple[Path, Optional[int], float]] = []
    for candidate in sorted(source_dir.rglob("*")):
        if not candidate.is_file():
            continue
        stem, version = split_version_suffix(candidate.stem)
        # Preserve the part before the extension; multi-dot names are fine.
        stem_normalized = normalize_name(stem.replace("_", " "))
        if stem_normalized.endswith(target_tail):
            if not prefer_newest:
                return candidate
            matches.append((candidate, version, candidate.stat().st_mtime))
    if not matches:
        return None
    return max(
        matches,
        key=lambda item: (
            item[2],
            -1 if item[1] is None else item[1],
            str(item[0]).lower(),
        ),
    )[0]


def unique_destination(path: Path) -> Path:
    """Return a unique path by appending a numeric suffix if required."""
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    counter = 1
    while True:
        candidate = path.with_name(f"{stem}_{counter}{suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def build_new_name(prefix: str, name: str, surname: str, suffix: str) -> str:
    """Construct the new filename from the prefix and reviewer name."""
    return build_base_name(prefix, name, surname) + suffix


def build_base_name(prefix: str, name: str, surname: str) -> str:
    """Construct the base filename without a file extension."""
    components = [prefix]
    for value in (name, surname):
        components.append(sanitized_component(value))
    return "_".join(components)


def unique_stem(dest_dir: Path, stem: str, suffixes: Sequence[str]) -> str:
    """Return a unique stem such that all suffix variations are free."""
    counter = 0
    while True:
        candidate = stem if counter == 0 else f"{stem}_{counter}"
        if all(not (dest_dir / f"{candidate}{suffix}").exists() for suffix in suffixes):
            return candidate
        counter += 1


def extract_docx_text(path: Path) -> str:
    """Extract text from a docx file using python-docx."""
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - depends on optional dependency
        raise RuntimeError(
            "python-docx is required for .docx conversion. Install with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)
    return "\n".join(lines)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    convert_docx = not args.no_docx_conversion

    source_dir: Optional[Path] = args.source_dir
    dest_dir: Path = args.dest_dir
    dry_run: bool = args.dry_run

    csv_parent: Optional[Path] = None
    try:
        with open_csv_source(args.csv_path) as (handle, _csv_label, csv_parent):
            if source_dir is None:
                source_dir = csv_parent or Path.cwd()

            if not source_dir.exists():
                print(f"Source directory not found: {source_dir}", file=sys.stderr)
                return 1
            dest_dir.mkdir(parents=True, exist_ok=True)

            missing: list[Tuple[str, str, str]] = []
            processed = 0
            processed_reviewers: set[str] = set()

            reader = csv.DictReader(handle)
            for row in reader:
                name = row.get("Name", "").strip()
                surname = row.get("Surname", "").strip()

                if not name and not surname:
                    if not args.skip_missing:
                        missing.append(("", "", "Missing Name and Surname in CSV row."))
                    continue

                reviewer_key = normalize_name(f"{name} {surname}")
                if args.prefer_newest and reviewer_key in processed_reviewers:
                    continue

                located = find_submission_file(
                    name,
                    surname,
                    source_dir,
                    prefer_newest=args.prefer_newest,
                )
                if located is None:
                    if not args.skip_missing:
                        missing.append((name, surname, "Matching file not found."))
                    continue

                suffix = located.suffix or ".txt"
                if suffix.lower() == ".docx" and convert_docx:
                    base = build_base_name(args.prefix, name, surname)
                    if args.prefer_newest:
                        txt_destination = dest_dir / f"{base}.txt"
                        docx_destination = dest_dir / f"{base}{suffix}"
                    else:
                        stem = unique_stem(dest_dir, base, [".txt", suffix])
                        txt_destination = dest_dir / f"{stem}.txt"
                        docx_destination = dest_dir / f"{stem}{suffix}"

                    if dry_run:
                        print(f"[DRY-RUN] {located} -> {txt_destination} (converted)")
                        print(f"[DRY-RUN] {located} -> {docx_destination}")
                    else:
                        try:
                            content = extract_docx_text(located)
                        except RuntimeError as exc:
                            if not args.skip_missing:
                                missing.append((name, surname, str(exc)))
                            continue
                        txt_destination.parent.mkdir(parents=True, exist_ok=True)
                        txt_destination.write_text(content, encoding="utf-8")
                        shutil.copy2(located, docx_destination)
                        print(f"Converted: {located} -> {txt_destination}")
                        print(f"Copied: {located} -> {docx_destination}")
                else:
                    new_name = build_new_name(args.prefix, name, surname, suffix)
                    destination = (
                        dest_dir / new_name
                        if args.prefer_newest
                        else unique_destination(dest_dir / new_name)
                    )

                    if dry_run:
                        print(f"[DRY-RUN] {located} -> {destination}")
                    else:
                        destination.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(located, destination)
                        print(f"Copied: {located} -> {destination}")
                processed += 1
                if args.prefer_newest:
                    processed_reviewers.add(reviewer_key)
    except FileNotFoundError:
        print(f"CSV not found: {args.csv_path}", file=sys.stderr)
        return 1
    except RuntimeError as exc:
        print(str(exc), file=sys.stderr)
        return 1

    if missing:
        print("\nUnable to process:", file=sys.stderr)
        for name, surname, reason in missing:
            label = (f"{name} {surname}".strip() or "Unknown reviewer").strip()
            if reason:
                print(f"  {label}: {reason}", file=sys.stderr)
            else:
                print(f"  {label}", file=sys.stderr)
        return 2 if not dry_run else 0

    if processed == 0:
        print("No submissions processed.", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
