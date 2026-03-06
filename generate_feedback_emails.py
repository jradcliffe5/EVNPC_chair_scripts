#!/usr/bin/env python3
"""Generate draft EVN PC feedback emails from a completed assessment file.

Reads the post-PC consensus assessment file, the code-mapping file, and the
reviewer-assignments CSV.  For each proposal it produces a draft email in a
single output .docx document (one proposal per page).

Optional features
-----------------
--split-tex   Write per-proposal LaTeX files (named <legacycode>_<evncode>.tex)
              from a directory of individual reviewer text files.  Reviewer
              names are replaced by "Reviewer 1", "Reviewer 2", etc. and
              internal (#) comment lines are removed.
--tex-only    Only produce the LaTeX files (requires --split-tex); skip
              writing the .docx output entirely.
--suffix-file Append the contents of a plain-text file (e.g. the grading-scale
              boilerplate) to every email.
"""

from __future__ import annotations

import argparse
import csv
import io
import math
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from textwrap import dedent
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI
import docx.enum.text as WD_ALIGN

# ---------------------------------------------------------------------------
# Shared constants (mirror reviews_to_latex.py where applicable)
# ---------------------------------------------------------------------------

SEPARATOR = "=" * 100
PROPOSAL_CODE_RE = re.compile(r"^[A-Z]\d{2}[A-Z]\d{3}$", re.IGNORECASE)

CONSENSUS_FIELD_NAMES = [
    "Grade",
    "Referee comments",
    "Strengths",
    "Weaknesses",
    "Technical review",
    "Time recommended",
]
CONSENSUS_FIELD_ALIASES = {name.lower(): name for name in CONSENSUS_FIELD_NAMES}

INDIVIDUAL_FIELD_NAMES = [
    "Grade",
    "General remark",
    "Strengths",
    "Weaknesses",
    "Referee comments",
    "Technical review",
    "Time recommended",
]
INDIVIDUAL_FIELD_ALIASES = {
    "grade": "Grade",
    "general remark": "General remark",
    "general remarks": "General remark",
    "strengths": "Strengths",
    "weaknesses": "Weaknesses",
    "referee comments": "Referee comments",
    "technical review": "Technical review",
    "time recommended": "Time recommended",
}


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ConsensusSummary:
    code: str          # EVN session code e.g. E26A001
    legacy_code: str   # Legacy code e.g. EC107
    pi: str
    title: str
    networks: str
    wavelengths: str
    grade: str = ""
    referee_comments: str = ""
    strengths: str = ""
    weaknesses: str = ""
    technical_review: str = ""
    time_recommended: str = ""


@dataclass
class IndividualReview:
    reviewer_label: str   # "Reviewer 1", "Reviewer 2", …
    grade: str = ""
    general_remark: str = ""
    referee_comments: str = ""
    strengths: str = ""
    weaknesses: str = ""
    technical_review: str = ""
    time_recommended: str = ""


# ---------------------------------------------------------------------------
# Argument parsing
# ---------------------------------------------------------------------------

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate draft EVN PC feedback emails as a single .docx.",
        epilog=dedent("""\
            Example:
              generate_feedback_emails.py \\
                -a EVNPC_2026A_assessment.txt \\
                -m evn_code_mapping.txt \\
                -r reviewer_assignments.txt \\
                -o feedback_draft.docx \\
                --split-tex tex_output/ \\
                --reviews-dir all_pc_reviews/ \\
                --suffix-file evn_pc_suffix_content.txt
        """),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "-a", "--assessment",
        type=Path,
        required=True,
        help="Completed consensus assessment file (post-PC, with grades and comments filled in).",
    )
    parser.add_argument(
        "-m", "--code-mapping",
        type=Path,
        default=None,
        help="File mapping EVN session codes to legacy proposal codes (e.g. E26A001 -> EC107).",
    )
    parser.add_argument(
        "-r", "--reviewer-assignments",
        type=Path,
        default=None,
        help="CSV file listing primary/secondary/additional reviewers per proposal.",
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=Path("feedback_draft.docx"),
        help="Output .docx file (default: feedback_draft.docx).",
    )
    parser.add_argument(
        "--suffix-file",
        type=Path,
        default=None,
        help="Plain-text file whose contents are appended to every email.",
    )
    parser.add_argument(
        "--split-tex",
        type=Path,
        default=None,
        metavar="DIR",
        help=(
            "Directory in which to write per-proposal LaTeX files "
            "(<legacycode>_<evncode>.tex).  Requires --reviews-dir."
        ),
    )
    parser.add_argument(
        "--tex-only",
        action="store_true",
        default=False,
        help=(
            "Only produce the LaTeX files (requires --split-tex); "
            "skip writing the .docx output entirely."
        ),
    )
    parser.add_argument(
        "--reviews-dir",
        type=Path,
        default=None,
        help="Directory of individual reviewer .txt files (used with --split-tex).",
    )
    parser.add_argument(
        "--session",
        default="",
        help="Session name displayed in the TOC title (e.g. 2026A).",
    )
    return parser.parse_args(argv)


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="cp1252")


def strip_internal_comments(text: str) -> str:
    """Remove lines whose first non-whitespace character is '#'."""
    lines = text.splitlines()
    kept = [line for line in lines if not line.strip().startswith("#")]
    return "\n".join(kept).strip()


def normalise_paragraphs(text: str) -> str:
    """Collapse soft line-wraps within paragraphs; preserve blank-line breaks."""
    if not text:
        return text
    paragraphs = re.split(r"\n\s*\n", text)
    cleaned = []
    for para in paragraphs:
        pieces = [l.strip() for l in para.splitlines() if l.strip()]
        if pieces:
            cleaned.append(" ".join(pieces))
    return "\n\n".join(cleaned)


def parse_header_columns(header: str) -> Tuple[str, str, str, str]:
    """Return (code, pi, networks, wavelengths) from the header line."""
    stripped = header.strip()
    columns = re.split(r"\s{2,}", stripped)
    if len(columns) >= 4 and PROPOSAL_CODE_RE.match(columns[0].strip()):
        code = columns[0].strip()
        pi = columns[1].strip()
        networks = " ".join(c.strip() for c in columns[2:-1] if c.strip())
        wavelengths = columns[-1].strip()
        return code, pi, networks, wavelengths
    # Fixed-width fallback
    return header[0:22].strip(), header[22:45].strip(), header[45:72].strip(), header[72:].strip()


def parse_value_block(
    lines: List[str],
    start_index: int,
    field_aliases: Dict[str, str],
) -> Tuple[str, str, int]:
    """Return (canonical_label, value_text, next_index)."""
    raw_line = lines[start_index]
    label_raw, rest = raw_line.split(":", 1)
    label_key = label_raw.strip().lower()
    canonical = field_aliases.get(label_key, label_raw.strip())

    paragraphs: List[str] = []
    current: List[str] = []
    first = rest.strip()
    if first:
        current.append(first)
    index = start_index + 1
    while index < len(lines):
        candidate = lines[index]
        stripped = candidate.strip()
        if stripped and ":" in stripped:
            candidate_label = stripped.split(":", 1)[0].strip().lower()
            if candidate_label in field_aliases:
                break
        if stripped == "":
            if current:
                paragraphs.append(" ".join(current).strip())
                current = []
            index += 1
            continue
        current.append(stripped)
        index += 1
    if current:
        paragraphs.append(" ".join(current).strip())
    return canonical, "\n\n".join(paragraphs).strip(), index


# ---------------------------------------------------------------------------
# Code mapping
# ---------------------------------------------------------------------------

def load_code_mapping(path: Optional[Path]) -> Dict[str, str]:
    """Return {EVN_code_upper: legacy_code}.

    Scans the full file for all occurrences of a proposal code (e.g. E26A001)
    followed by a legacy code token.  This handles files where multiple
    proposals appear on a single line (a common formatting artefact).
    """
    if path is None or not path.is_file():
        return {}
    text = read_text(path)
    # Match an EVN session code followed by a legacy code (e.g. EC107, GS047)
    pattern = re.compile(
        r"\b([A-Z]\d{2}[A-Z]\d{3})\s+([A-Z]{1,2}\d{3,5})\b", re.IGNORECASE
    )
    mapping: Dict[str, str] = {}
    for m in pattern.finditer(text):
        evn_code = m.group(1).upper()
        legacy_code = m.group(2)
        if evn_code not in mapping:
            mapping[evn_code] = legacy_code
    return mapping


# ---------------------------------------------------------------------------
# Reviewer assignments
# ---------------------------------------------------------------------------

def load_reviewer_assignments(path: Optional[Path]) -> Dict[str, List[str]]:
    """Return {proposal_code: [primary, secondary, additional1, ...]}."""
    if path is None or not path.is_file():
        return {}
    data = read_text(path)
    # Read only up to a blank line after the header (the CSV section)
    csv_lines = []
    for line in data.splitlines():
        if not line.strip():
            if csv_lines:
                break
        else:
            csv_lines.append(line)
    if not csv_lines:
        return {}
    reader = csv.DictReader(io.StringIO("\n".join(csv_lines)))
    assignments: Dict[str, List[str]] = {}
    for row in reader:
        row_l = {k.lower().strip(): (v or "").strip() for k, v in row.items()}
        code = row_l.get("proposal", "")
        if not code:
            continue
        ordered: List[str] = []
        for col in ("primary reviewer", "secondary reviewer"):
            name = row_l.get(col, "")
            if name:
                ordered.append(name)
        # Additional reviewers: any column that starts with "additional"
        for key in row_l:
            if key.startswith("additional"):
                name = row_l[key]
                if name:
                    ordered.append(name)
        assignments[code.upper()] = ordered
    return assignments


# ---------------------------------------------------------------------------
# Parse consensus assessment file
# ---------------------------------------------------------------------------

def parse_assessment_file(
    path: Path,
    code_mapping: Dict[str, str],
) -> List[ConsensusSummary]:
    """Parse the post-PC assessment file into ConsensusSummary objects."""
    content = read_text(path)
    blocks = [b.strip("\n") for b in content.split(SEPARATOR) if b.strip()]
    summaries: List[ConsensusSummary] = []
    for block in blocks:
        lines = [l.rstrip() for l in block.splitlines()]
        if not lines or not lines[0].strip():
            continue
        code, pi, networks, wavelengths = parse_header_columns(lines[0])
        title = lines[1].strip() if len(lines) > 1 else ""
        if not PROPOSAL_CODE_RE.match(code):
            continue

        fields = {name: "" for name in CONSENSUS_FIELD_NAMES}
        index = 2
        while index < len(lines):
            stripped = lines[index].strip()
            if not stripped:
                index += 1
                continue
            if ":" not in stripped:
                index += 1
                continue
            label_key = stripped.split(":", 1)[0].strip().lower()
            # Skip "Primary reviewer:" / "Secondary reviewer:" lines
            if label_key in ("primary reviewer", "secondary reviewer"):
                index += 1
                continue
            if label_key in CONSENSUS_FIELD_ALIASES:
                canonical, value, index = parse_value_block(
                    lines, index, CONSENSUS_FIELD_ALIASES
                )
                fields[canonical] = value
            else:
                index += 1

        legacy = code_mapping.get(code.upper(), "")
        summaries.append(ConsensusSummary(
            code=code.upper(),
            legacy_code=legacy,
            pi=pi,
            title=title,
            networks=networks,
            wavelengths=wavelengths,
            grade=fields["Grade"],
            referee_comments=strip_internal_comments(fields["Referee comments"]),
            strengths=strip_internal_comments(fields["Strengths"]),
            weaknesses=strip_internal_comments(fields["Weaknesses"]),
            technical_review=strip_internal_comments(fields["Technical review"]),
            time_recommended=fields["Time recommended"],
        ))
    return summaries


# ---------------------------------------------------------------------------
# Parse individual reviews for --split-tex
# ---------------------------------------------------------------------------

def parse_individual_review_file(
    path: Path,
    assignments: Dict[str, List[str]],
) -> Dict[str, List[IndividualReview]]:
    """Parse one reviewer's file; return {proposal_code: [IndividualReview]}."""
    content = read_text(path)
    blocks = [b.strip("\n") for b in content.split(SEPARATOR) if b.strip()]
    result: Dict[str, List[IndividualReview]] = {}

    # Infer reviewer name from filename (prefix_First_Last.txt)
    stem = path.stem
    parts = stem.split("_", 1)
    reviewer_name = parts[1].replace("_", " ").strip() if len(parts) == 2 else stem

    for block in blocks:
        lines = [l.rstrip() for l in block.splitlines()]
        if not lines or not lines[0].strip():
            continue
        code, _, _, _ = parse_header_columns(lines[0])
        if not PROPOSAL_CODE_RE.match(code):
            continue
        code = code.upper()

        # Determine reviewer label for this proposal
        ordered = assignments.get(code, [])
        label = "Reviewer ?"
        for idx, name in enumerate(ordered, start=1):
            if name.lower() == reviewer_name.lower():
                label = f"Reviewer {idx}"
                break

        fields: Dict[str, str] = {name: "" for name in INDIVIDUAL_FIELD_NAMES}
        index = 2
        while index < len(lines):
            stripped = lines[index].strip()
            if not stripped:
                index += 1
                continue
            if ":" not in stripped:
                index += 1
                continue
            label_key = stripped.split(":", 1)[0].strip().lower()
            if label_key in ("primary reviewer", "secondary reviewer"):
                index += 1
                continue
            if label_key in INDIVIDUAL_FIELD_ALIASES:
                canonical, value, index = parse_value_block(
                    lines, index, INDIVIDUAL_FIELD_ALIASES
                )
                fields[canonical] = value
            else:
                index += 1

        if not any(fields[n] for n in INDIVIDUAL_FIELD_NAMES):
            continue

        review = IndividualReview(
            reviewer_label=label,
            grade=fields["Grade"],
            general_remark=strip_internal_comments(fields["General remark"]),
            referee_comments=strip_internal_comments(fields["Referee comments"]),
            strengths=strip_internal_comments(fields["Strengths"]),
            weaknesses=strip_internal_comments(fields["Weaknesses"]),
            technical_review=strip_internal_comments(fields["Technical review"]),
            time_recommended=fields["Time recommended"],
        )
        result.setdefault(code, []).append(review)
    return result


def load_all_individual_reviews(
    reviews_dir: Path,
    assignments: Dict[str, List[str]],
) -> Dict[str, List[IndividualReview]]:
    """Merge all per-reviewer files into {proposal_code: [IndividualReview]}."""
    combined: Dict[str, List[IndividualReview]] = {}
    for fp in sorted(reviews_dir.glob("*.txt")):
        per_file = parse_individual_review_file(fp, assignments)
        for code, reviews in per_file.items():
            combined.setdefault(code, []).extend(reviews)
    # Sort each proposal's reviews by reviewer label number
    def label_key(r: IndividualReview) -> int:
        m = re.search(r"\d+", r.reviewer_label)
        return int(m.group()) if m else 999
    for code in combined:
        combined[code].sort(key=label_key)
    return combined


# ---------------------------------------------------------------------------
# LaTeX helpers (for --split-tex)
# ---------------------------------------------------------------------------

UNICODE_REPLACEMENTS = {
    "\u00a0": " ", "\u202f": " ",
    "\u2010": "-", "\u2011": "-", "\u2012": "-",
    "\u2013": "--", "\u2014": "---", "\u2212": "-",
    "\u2026": "...",
    "\u2018": "'", "\u2019": "'",
    "\u201c": "``", "\u201d": "''",
    "\u2032": "'", "\u2033": r"$''$",
    "\u00b5": r"$\mu$", "\u00b0": r"$^{\circ}$",
    "\u00d7": r"$\times$", "\u223c": r"$\sim$",
    "\u2264": r"$\le$", "\u2265": r"$\ge$",
    "\u03b1": r"$\alpha$", "\u03b2": r"$\beta$",
    "\u03b3": r"$\gamma$", "\u03c0": r"$\pi$",
    "\u03bc": r"$\mu$",
}

UNICODE_LATEX_DECLARATIONS = {
    "00B1": r"\ensuremath{\pm}",
    "0144": r"\'{n}",
    "03C3": r"\ensuremath{\sigma}",
    "2192": r"\ensuremath{\rightarrow}",
    "2248": r"\ensuremath{\approx}",
    "FF5E": r"\textasciitilde{}",
}


def latex_escape(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    for src, repl in UNICODE_REPLACEMENTS.items():
        text = text.replace(src, repl)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    pattern = re.compile("|".join(re.escape(k) for k in replacements))
    return pattern.sub(lambda m: replacements[m.group()], text)


def tex_paragraphs(text: str) -> str:
    """Escape text and separate paragraphs with blank lines for LaTeX."""
    if not text:
        return ""
    paras = re.split(r"\n\s*\n", text)
    result = []
    for para in paras:
        pieces = [l.strip() for l in para.splitlines() if l.strip()]
        if pieces:
            result.append(latex_escape(" ".join(pieces)))
    return "\n\n".join(result)


def build_proposal_tex(
    summary: ConsensusSummary,
    individual_reviews: List[IndividualReview],
) -> str:
    """Build a complete standalone LaTeX document for one proposal."""
    declarations = "\n".join(
        rf"\DeclareUnicodeCharacter{{{cp}}}{{{lt}}}"
        for cp, lt in UNICODE_LATEX_DECLARATIONS.items()
    )
    code_label = f"{summary.code}/{summary.legacy_code}" if summary.legacy_code else summary.code
    section_title = f"{code_label}: {summary.title}" if summary.title else code_label

    lines: List[str] = [
        r"\documentclass[a4paper,11pt]{article}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{lmodern}",
        r"\usepackage[margin=2.5cm]{geometry}",
        r"\usepackage[hidelinks]{hyperref}",
        r"\usepackage{xcolor}",
        declarations,
        r"\setlength{\parindent}{0pt}",
        r"\setlength{\parskip}{6pt}",
        r"\begin{document}",
        rf"\section*{{{latex_escape(section_title)}}}",
    ]
    if summary.pi:
        lines.append(rf"\textbf{{PI:}} {latex_escape(summary.pi)}\\")
    if summary.networks:
        lines.append(rf"\textbf{{Networks:}} {latex_escape(summary.networks)}\\")
    if summary.wavelengths:
        lines.append(rf"\textbf{{Requested wavelengths:}} {latex_escape(summary.wavelengths)}\\")

    lines.append("")
    for review in individual_reviews:
        lines.append(rf"\subsection*{{{latex_escape(review.reviewer_label)}}}")
        if review.grade:
            lines.append(rf"\textbf{{Grade:}} {latex_escape(review.grade)}\\")
        if review.time_recommended:
            lines.append(rf"\textbf{{Time recommended:}} {latex_escape(review.time_recommended)}\\")

        for label, attr in (
            ("General remark", "general_remark"),
            ("Referee comments", "referee_comments"),
            ("Strengths", "strengths"),
            ("Weaknesses", "weaknesses"),
        ):
            value = getattr(review, attr)
            if value:
                lines.append(rf"\textbf{{{label}}}\\")
                lines.append(tex_paragraphs(value))
                lines.append("")

        if review.technical_review:
            lines.append(r"\textbf{Technical review}\\")
            lines.append(tex_paragraphs(review.technical_review))
            lines.append("")

    lines.append(r"\end{document}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Word comment helpers
# ---------------------------------------------------------------------------

_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_WNS}}}"
_COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"


def _get_or_create_comments_part(doc: Document):
    dp = doc.part
    for rel in dp.rels.values():
        if rel.reltype == _COMMENTS_REL:
            return rel.target_part
    root = etree.fromstring(f'<w:comments xmlns:w="{_WNS}"/>'.encode())
    cp = XmlPart(PackURI("/word/comments.xml"), _COMMENTS_CT, root, dp.package)
    dp.relate_to(cp, _COMMENTS_REL)
    return cp


def add_word_comment(doc: Document, para, comment_text: str, cid: int,
                     author: str = "EVN PC") -> None:
    """Attach a Word comment annotation to *para* with the given text."""
    cp = _get_or_create_comments_part(doc)
    root = cp._element

    # Build <w:comment> element
    comment_el = etree.SubElement(root, f"{_W}comment")
    comment_el.set(f"{_W}id", str(cid))
    comment_el.set(f"{_W}author", author)
    comment_el.set(f"{_W}date", "2026-01-01T00:00:00Z")
    initials = "".join(w[0].upper() for w in author.split() if w)
    comment_el.set(f"{_W}initials", initials)
    cp_p = etree.SubElement(comment_el, f"{_W}p")
    cp_r = etree.SubElement(cp_p, f"{_W}r")
    cp_t = etree.SubElement(cp_r, f"{_W}t")
    cp_t.text = comment_text

    # Add commentRangeStart before the first run in the paragraph
    p_xml = para._p
    crs = etree.Element(f"{_W}commentRangeStart")
    crs.set(f"{_W}id", str(cid))
    first_run = p_xml.find(f"{_W}r")
    if first_run is not None:
        p_xml.insert(list(p_xml).index(first_run), crs)
    else:
        p_xml.append(crs)

    # Add commentRangeEnd and commentReference run at end of paragraph
    cre = etree.Element(f"{_W}commentRangeEnd")
    cre.set(f"{_W}id", str(cid))
    p_xml.append(cre)
    ref_r = etree.SubElement(p_xml, f"{_W}r")
    ref_rpr = etree.SubElement(ref_r, f"{_W}rPr")
    ref_style = etree.SubElement(ref_rpr, f"{_W}rStyle")
    ref_style.set(f"{_W}val", "CommentReference")
    ref_ref = etree.SubElement(ref_r, f"{_W}commentReference")
    ref_ref.set(f"{_W}id", str(cid))


# ---------------------------------------------------------------------------
# Build the .docx document
# ---------------------------------------------------------------------------

def add_toc_page(doc: Document, session: str) -> None:
    """Insert a title and Word built-in TOC field.

    Word will render this as a linked, page-numbered table of contents when
    the document is opened (fields are marked dirty so they update on open).
    The TOC collects all Heading 1 paragraphs — one is added at the top of
    every proposal page.
    """
    title = f"EVN PC {session} Feedback" if session else "EVN PC Feedback"
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(18)

    doc.add_paragraph()

    # Build the TOC field: fldChar(begin) + instrText + fldChar(separate) +
    # placeholder text + fldChar(end).  w:dirty="true" tells Word to update
    # the field the first time the document is opened.
    p = doc.add_paragraph()
    p_xml = p._p

    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.set(qn("w:dirty"), "true")
    r_begin.append(fc_begin)
    p_xml.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-1" \\h \\z '
    r_instr.append(instr)
    p_xml.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    p_xml.append(r_sep)

    r_ph = OxmlElement("w:r")
    t_ph = OxmlElement("w:t")
    t_ph.text = (
        "Right-click this line and select \u2018Update Field\u2019 "
        "to generate the table of contents."
    )
    r_ph.append(t_ph)
    p_xml.append(r_ph)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    p_xml.append(r_end)


def add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def set_font(run, bold: bool = False, size_pt: Optional[int] = None,
             italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def add_heading_para(doc: Document, text: str, size_pt: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


def add_field_para(doc: Document, label: str, value: str) -> None:
    """Add a paragraph like 'Label: value' with the label bolded."""
    if not value:
        return
    p = doc.add_paragraph()
    r_label = p.add_run(f"{label}: ")
    r_label.bold = True
    p.add_run(value)


_INLINE_MARKUP_RE = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)", re.DOTALL)


def _add_formatted_runs(para, text: str) -> None:
    """Add runs to *para* honouring **bold**, *italic*, and ***bold+italic*** markup."""
    last = 0
    for m in _INLINE_MARKUP_RE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(2) is not None:        # ***bold+italic***
            r = para.add_run(m.group(2))
            r.bold = True
            r.italic = True
        elif m.group(3) is not None:      # **bold**
            para.add_run(m.group(3)).bold = True
        elif m.group(4) is not None:      # *italic*
            para.add_run(m.group(4)).italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def add_body_text(doc: Document, text: str) -> None:
    """Add multi-paragraph body text, splitting on blank lines.

    Inline markup supported in the source text:
        **bold**   *italic*   ***bold+italic***
    """
    if not text:
        return
    paras = re.split(r"\n\s*\n", text)
    for para in paras:
        pieces = [l.strip() for l in para.splitlines() if l.strip()]
        if pieces:
            p = doc.add_paragraph()
            _add_formatted_runs(p, " ".join(pieces))


def build_email_docx(
    summaries: List[ConsensusSummary],
    suffix_text: str,
    output_path: Path,
    assignments: Optional[Dict[str, List[str]]] = None,
    session: str = "",
) -> None:
    doc = Document()
    # Use a clean default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tell Word to update all fields (including the TOC) when opening the file
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)

    # --- Table of contents page ---
    add_toc_page(doc, session)
    add_page_break(doc)

    comment_id = 0

    for idx, s in enumerate(summaries):
        if idx > 0:
            add_page_break(doc)

        code_label = f"{s.legacy_code}/{s.code}" if s.legacy_code else s.code

        # Heading 1 paragraph — picked up by the Word TOC field
        heading_text = f"{code_label} \u2014 {s.pi}"
        doc.add_heading(heading_text, level=1)

        doc.add_paragraph()

        # --- Greeting ---
        doc.add_paragraph("Dear colleague,")

        # --- Opening paragraph with placeholders ---
        p = doc.add_paragraph()
        r_intro = p.add_run("[I am pleased to inform you that / I regret to inform you that] ")
        r_intro.bold = True
        p.add_run("your proposal ")
        p.add_run(s.code).bold = True
        p.add_run(f" (legacy code {s.legacy_code}) " if s.legacy_code else " ")

        # "was/was not approved" — bold placeholder
        r = p.add_run("[was / was not] approved")
        r.bold = True

        p.add_run(" by the EVN Programme Committee (PC). ")

        r_outcome = p.add_run(
            "[This proposal will be put in the observing queue. / "
            "The EVN PC strongly encourages the proposal team to resubmit a revised "
            "version of their proposal, taking into account the comments presented here.]"
        )
        r_outcome.bold = True

        doc.add_paragraph(
            "Please acknowledge the receipt of this email and forward copies of this "
            "message to your co-proposers."
        )
        doc.add_paragraph(
            "Below you can find the consensus report from EVN PC and an explanation of "
            "the grading scheme. A copy of the summary comments and rating given, "
            "including the individual comments of the PC members are in the attached "
            "pdf document.\n\n"
            "The User Support Group of the Joint Institute for VLBI ERIC (JIVE) can "
            "provide assistance with proposal submission, schedule preparation, and "
            "data analysis \u2013 please contact usersupport@jive.eu."
        )

        doc.add_paragraph()

        # --- Proposal details table ---
        p = doc.add_paragraph()
        p.add_run("Code:    ").bold = True
        p.add_run(code_label)
        p.add_run("    PI:   ").bold = True
        p.add_run(s.pi)

        # Word comment: Primary / Secondary reviewer
        if assignments is not None:
            reviewers = assignments.get(s.code, [])
            primary = reviewers[0] if len(reviewers) > 0 else "?"
            secondary = reviewers[1] if len(reviewers) > 1 else "?"
            add_word_comment(
                doc, p,
                f"Primary reviewer: {primary} | Secondary reviewer: {secondary}",
                cid=comment_id,
            )
            comment_id += 1

        p2 = doc.add_paragraph()
        p2.add_run("Title:  ").bold = True
        p2.add_run(s.title)

        doc.add_paragraph()

        # --- EVN PC Rating ---
        p3 = doc.add_paragraph()
        p3.add_run("EVN PC Rating: ").bold = True
        p3.add_run(s.grade)

        p4 = doc.add_paragraph()
        p4.add_run("Time: ").bold = True
        p4.add_run(s.time_recommended)

        doc.add_paragraph()

        # --- Summary of comments ---
        p_sum = doc.add_paragraph()
        p_sum.add_run("Summary of comments:").bold = True
        doc.add_paragraph()

        if s.referee_comments:
            add_body_text(doc, s.referee_comments)
        else:
            p_ph = doc.add_paragraph()
            r_ph = p_ph.add_run("[PC members to include summarised comments here]")
            r_ph.bold = True
            r_ph.italic = True

        if s.strengths:
            p_s = doc.add_paragraph()
            p_s.add_run("Strengths: ").bold = True
            add_body_text(doc, s.strengths)

        if s.weaknesses:
            p_w = doc.add_paragraph()
            p_w.add_run("Weaknesses: ").bold = True
            add_body_text(doc, s.weaknesses)

        doc.add_paragraph()

        doc.add_paragraph(
            "Please read the comments from the individual reviewers carefully, but note "
            "that these were made before the EVN PC discussions. The summary above is "
            "the definitive evaluation of the proposal."
        )

        doc.add_paragraph()

        # --- Technical review ---
        p5 = doc.add_paragraph()
        p5.add_run("Technical review: ").bold = True
        if s.technical_review:
            add_body_text(doc, s.technical_review)

        # --- Suffix ---
        if suffix_text:
            doc.add_paragraph()
            doc.add_paragraph("\u2014" * 40)  # horizontal rule as dashes
            add_body_text(doc, suffix_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote feedback email draft to {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    # Load code mapping
    code_mapping = load_code_mapping(args.code_mapping)

    # Load reviewer assignments
    assignments = load_reviewer_assignments(args.reviewer_assignments)

    # Parse consensus assessment
    if not args.assessment.is_file():
        print(f"Assessment file not found: {args.assessment}", file=sys.stderr)
        return 1
    summaries = parse_assessment_file(args.assessment, code_mapping)
    if not summaries:
        print("No proposals parsed from assessment file.", file=sys.stderr)
        return 1
    print(f"Parsed {len(summaries)} proposals from {args.assessment}")

    # Validate --tex-only usage
    if args.tex_only and args.split_tex is None:
        print("--tex-only requires --split-tex.", file=sys.stderr)
        return 1

    # Load suffix
    suffix_text = ""
    if args.suffix_file:
        if not args.suffix_file.is_file():
            print(f"Suffix file not found: {args.suffix_file}", file=sys.stderr)
            return 1
        suffix_text = read_text(args.suffix_file).strip()

    # Build the .docx (unless --tex-only was requested)
    if not args.tex_only:
        build_email_docx(summaries, suffix_text, args.output,
                         assignments=assignments, session=args.session)

    # --split-tex: generate per-proposal LaTeX files
    if args.split_tex is not None:
        if args.reviews_dir is None or not args.reviews_dir.is_dir():
            print(
                "--split-tex requires --reviews-dir pointing to a valid directory.",
                file=sys.stderr,
            )
            return 1
        individual = load_all_individual_reviews(args.reviews_dir, assignments)
        args.split_tex.mkdir(parents=True, exist_ok=True)
        for s in summaries:
            reviews_for_proposal = individual.get(s.code, [])
            tex_content = build_proposal_tex(s, reviews_for_proposal)
            stem = f"{s.legacy_code}_{s.code}" if s.legacy_code else s.code
            out_tex = args.split_tex / f"{stem}.tex"
            out_tex.write_text(tex_content, encoding="utf-8")
            print(f"  Wrote {out_tex}")
        print(f"Wrote {len(summaries)} LaTeX files to {args.split_tex}/")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
