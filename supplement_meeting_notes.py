#!/usr/bin/env python3
"""
supplement_meeting_notes.py
----------------------------
Automatically supplements EVN PC meeting notes (DOCX) with content from:
  - Zoom closed captions (.txt)
  - Zoom chat transcript (.txt)
  - Audio recordings (.m4a / .mp3 / .wav) — transcribed via mlx-whisper

Usage:
    python3 supplement_meeting_notes.py <notes.docx> [options]

Options:
    --captions FILE     Closed caption text file
    --chat FILE         Chat log text file
    --audio FILE [FILE] One or more audio files (transcribed with mlx-whisper)
    --output FILE       Output DOCX path (default: <input>_supplemented.docx)
    --model MODEL       mlx-whisper model (default: mlx-community/whisper-large-v3-turbo)
    --api-key KEY       Anthropic API key (or set ANTHROPIC_API_KEY env var)
    --no-ai             Skip Claude AI synthesis; just append raw transcripts as notes

Requirements:
    pip install python-docx mlx-whisper anthropic

Example:
    python3 supplement_meeting_notes.py EVNPC_2026A_agenda.docx \\
        --captions closed_caption.txt \\
        --chat chat.txt \\
        --audio audio1.m4a audio2.m4a audio3.m4a
"""

import argparse
import subprocess
import sys
import json
import re
from pathlib import Path


# ── helpers ──────────────────────────────────────────────────────────────────

def read_text_file(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def transcribe_audio(audio_path, model="mlx-community/whisper-large-v3-turbo"):
    """Transcribe an audio file using mlx-whisper. Returns timestamped transcript string."""
    try:
        import mlx_whisper
    except ImportError:
        print("  mlx_whisper not installed. Run: pip install mlx-whisper", file=sys.stderr)
        return None

    print(f"  Transcribing {Path(audio_path).name} ...")
    result = mlx_whisper.transcribe(
        str(audio_path),
        path_or_hf_repo=model,
        word_timestamps=False,
    )
    lines = []
    for seg in result.get("segments", []):
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if text:
            lines.append(f"[{start:.1f}s] {text}")
    return "\n".join(lines)


def extract_docx_text(docx_path):
    """Return the plain text of a DOCX as a list of (para_index, text) tuples."""
    from docx import Document
    doc = Document(str(docx_path))
    return [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.text.strip()]


# ── Claude API synthesis ──────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert scientific meeting secretary helping to supplement EVN (European VLBI Network) Programme Committee meeting notes.

You will receive:
1. The existing meeting notes (from a DOCX agenda/notes document)
2. Supplementary sources: closed captions, chat log, and/or audio transcripts

Your task: For each section of the meeting notes, identify important details, decisions, discussions, or context that appear in the supplementary sources but are NOT already captured in the notes. Produce a structured JSON list of insertions.

Rules:
- Only add genuinely new information not already in the notes
- Be concise but complete — include specific names, numbers, decisions, action items
- Preserve technical terms (telescope names, proposal codes, grades, frequencies)
- Note the source type (captions/chat/audio)
- Output JSON only, no preamble

Output format:
[
  {
    "anchor": "exact substring of the notes paragraph to insert after",
    "note": "The supplementary text to insert (1-3 sentences, starting with [Source])",
    "source": "captions|chat|audio"
  },
  ...
]
"""


def call_claude(notes_text, supplementary_sources):
    """Call Claude via the Claude Code CLI to synthesise supplementary notes."""
    import shutil
    if not shutil.which("claude"):
        print("  'claude' CLI not found. Ensure Claude Code is installed and on PATH.", file=sys.stderr)
        return []

    sources_block = ""
    for label, content in supplementary_sources.items():
        if content:
            truncated = content[:40000] if len(content) > 40000 else content
            sources_block += f"\n\n=== {label.upper()} ===\n{truncated}"

    prompt = f"""{SYSTEM_PROMPT}

MEETING NOTES:
{notes_text[:20000]}

SUPPLEMENTARY SOURCES:{sources_block}

Please produce the JSON insertion list as described."""

    print("  Calling Claude CLI for intelligent synthesis...")
    result = subprocess.run(
        ["claude", "-p"],
        input=prompt,
        capture_output=True,
        text=True,
        timeout=300,
    )

    if result.returncode != 0:
        print(f"  Claude CLI error: {result.stderr.strip()}", file=sys.stderr)
        return []

    response_text = result.stdout.strip()

    json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError as e:
            print(f"  Warning: Could not parse Claude JSON response: {e}", file=sys.stderr)
            print(f"  Raw response: {response_text[:500]}", file=sys.stderr)
            return []
    return []


# ── DOCX editing ─────────────────────────────────────────────────────────────

def make_note_para(note_text, source="captions"):
    """Create a formatted supplementary note paragraph element."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    color_map = {
        "captions": "1F4E79",   # dark blue
        "chat":     "7B2C8C",   # purple
        "audio":    "375623",   # dark green
    }
    icon_map = {
        "captions": "📝",
        "chat":     "💬",
        "audio":    "🎙",
    }
    color = color_map.get(source, "1F4E79")
    icon = icon_map.get(source, "📝")

    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    rpr.append(color_el)
    rpr.append(OxmlElement('w:i'))

    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '18')   # 9pt
    rpr.append(sz_el)
    sz_cs = OxmlElement('w:szCs')
    sz_cs.set(qn('w:val'), '18')
    rpr.append(sz_cs)

    r.append(rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = f"{icon} {note_text}"
    r.append(t)
    p.append(r)
    return p


def apply_insertions(doc, insertions):
    """Apply a list of {anchor, note, source} dicts to the document."""
    inserted_count = 0
    for item in insertions:
        anchor = item.get("anchor", "")
        note = item.get("note", "")
        source = item.get("source", "captions")

        if not anchor or not note:
            continue

        target_para = None
        for para in doc.paragraphs:
            if anchor in para.text:
                target_para = para
                break

        if target_para is None:
            print(f"  ⚠ Could not find anchor: '{anchor[:60]}'", file=sys.stderr)
            continue

        note_p = make_note_para(note, source)
        target_para._element.addnext(note_p)
        inserted_count += 1

    return inserted_count


def append_full_transcripts(doc, supplementary_sources):
    """Append full transcripts as a new section at end of document (fallback / no-AI mode)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    from docx import Document

    # Add a section break and heading
    doc.add_page_break()
    h = doc.add_heading("Supplementary Source Transcripts", level=1)

    for label, content in supplementary_sources.items():
        if not content:
            continue
        doc.add_heading(label, level=2)
        # Add in chunks to avoid huge paragraphs
        for chunk in content.split('\n'):
            if chunk.strip():
                p = doc.add_paragraph(chunk)
                p.runs[0].font.size = Pt(8)


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("notes", help="Input meeting notes DOCX file")
    parser.add_argument("--captions", help="Closed captions text file")
    parser.add_argument("--chat", help="Chat log text file")
    parser.add_argument("--audio", nargs="+", help="Audio files to transcribe")
    parser.add_argument("--output", help="Output DOCX path")
    parser.add_argument("--model", default="mlx-community/whisper-large-v3-turbo",
                        help="Whisper model for transcription")
    parser.add_argument("--no-ai", action="store_true",
                        help="Skip Claude AI; just append transcripts to document")
    args = parser.parse_args()

    notes_path = Path(args.notes)
    if not notes_path.exists():
        print(f"Error: notes file not found: {notes_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output) if args.output else notes_path.with_stem(notes_path.stem + "_supplemented")

    # ── gather supplementary sources ──
    supplementary_sources = {}

    if args.captions:
        print(f"Reading captions: {args.captions}")
        supplementary_sources["Closed Captions"] = read_text_file(args.captions)

    if args.chat:
        print(f"Reading chat: {args.chat}")
        supplementary_sources["Chat Log"] = read_text_file(args.chat)

    if args.audio:
        combined_transcripts = []
        for audio_file in args.audio:
            apath = Path(audio_file)
            # Check for cached transcript (both naming conventions: .transcript.txt and _transcript.txt)
            cache_path = apath.with_suffix(".transcript.txt")
            alt_cache_path = apath.parent / (apath.stem + "_transcript.txt")
            if cache_path.exists():
                print(f"Using cached transcript: {cache_path.name}")
                transcript = read_text_file(cache_path)
            elif alt_cache_path.exists():
                print(f"Using cached transcript: {alt_cache_path.name}")
                transcript = read_text_file(alt_cache_path)
            else:
                transcript = transcribe_audio(str(apath), model=args.model)
                if transcript:
                    cache_path.write_text(transcript, encoding="utf-8")
                    print(f"  Saved transcript cache: {cache_path.name}")
            if transcript:
                combined_transcripts.append(f"--- {apath.name} ---\n{transcript}")
        if combined_transcripts:
            supplementary_sources["Audio Transcripts"] = "\n\n".join(combined_transcripts)

    if not supplementary_sources:
        print("Warning: no supplementary sources provided. Nothing to add.", file=sys.stderr)
        sys.exit(0)

    # ── load document ──
    from docx import Document
    print(f"\nLoading notes document: {notes_path.name}")
    doc = Document(str(notes_path))
    notes_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # ── synthesise with Claude or append raw ──
    if args.no_ai:
        print("Appending full transcripts to document (no-AI mode)...")
        append_full_transcripts(doc, supplementary_sources)
        total_inserted = sum(len(v.splitlines()) for v in supplementary_sources.values())
        print(f"Appended {len(supplementary_sources)} source(s).")
    else:
        insertions = call_claude(notes_text, supplementary_sources)
        if insertions:
            print(f"  Got {len(insertions)} insertion suggestions from Claude.")
            n = apply_insertions(doc, insertions)
            print(f"  Successfully inserted {n} notes into document.")
        else:
            print("  No insertions returned by Claude. Falling back to appending transcripts.")
            append_full_transcripts(doc, supplementary_sources)

    # ── save ──
    doc.save(str(output_path))
    print(f"\n✓ Supplemented document saved to: {output_path}")


if __name__ == "__main__":
    main()
