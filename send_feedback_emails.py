#!/usr/bin/env python3
"""Send EVN PC feedback emails to PIs with the compiled PDF review attachment.

For each proposal in the completed feedback .docx the script:
  1. Compiles any unbuilt/stale .tex files in --pdf-dir using pdflatex.
  2. Extracts the email body text (plain text from docx paragraphs).
  3. Looks up the PI's email address from the pi-emails file produced by
     proposal_to_review_template.py (--pi-emails-file).
  4. Locates the compiled PDF in --pdf-dir (named <legacy>_<code>.pdf or
     <code>.pdf).
  5. Sends the email with the PDF attached.

Run with --dry-run (the default) to preview without sending.
Pass --send explicitly (via run-PC-scripts.sh) to deliver messages.
"""

from __future__ import annotations

import argparse
import html as _html
import imaplib
import json
import re
import shutil
import smtplib
import subprocess
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime, timezone
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Set

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PROPOSAL_RE = re.compile(r"[A-Z]\d{2}[A-Z]\d{3}", re.IGNORECASE)

# Matches the choice-placeholder style used in generate_feedback_emails.py,
# e.g. "[I am pleased … / I regret …]" or "[was / was not]".
UNRESOLVED_CHOICE_RE = re.compile(r"\[[^\]]+/[^\]]+\]")

# The placeholder injected when no referee comments are present yet.
FEEDBACK_PLACEHOLDER = "[PC members to include summarised comments here]"

DEFAULT_SUBJECT = "[Results from the EVN PC] {title} [{code_label}]"


# ---------------------------------------------------------------------------
# Data structure
# ---------------------------------------------------------------------------

@dataclass
class ProposalEmail:
    code: str
    legacy_code: str
    pi_name: str
    title: str = ""
    paragraphs: List[str] = field(default_factory=list)       # plain text
    html_paragraphs: List[str] = field(default_factory=list)  # HTML runs
    unresolved_placeholders: List[str] = field(default_factory=list)

    @property
    def body_text(self) -> str:
        # Separate every non-empty paragraph with a blank line — standard
        # plain-text email formatting and prevents sections running together.
        parts = [para.strip() for para in self.paragraphs if para.strip()]
        return "\n\n".join(parts)

    @property
    def html_body(self) -> str:
        parts = [p for p in self.html_paragraphs if p.strip()]
        inner = "\n".join(f"<p>{p}</p>" for p in parts)
        return (
            '<html><body style="font-family:Arial,sans-serif;font-size:14px;">\n'
            + inner
            + "\n</body></html>"
        )


# ---------------------------------------------------------------------------
# Argument parsing
# ---------------------------------------------------------------------------

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Send EVN PC feedback emails to PIs with attached PDF reviews.",
    )
    parser.add_argument(
        "--feedback-docx", "-d",
        type=Path,
        required=True,
        help="Completed feedback .docx produced by generate_feedback_emails.py.",
    )
    parser.add_argument(
        "--pi-emails-file", "-p",
        type=Path,
        required=True,
        help="File of 'CODE: email' lines produced by proposal_to_review_template.py.",
    )
    parser.add_argument(
        "--pdf-dir",
        type=Path,
        default=None,
        help=(
            "Directory containing compiled PDFs from the feedback_tex/ LaTeX files. "
            "Expected names: <legacy>_<code>.pdf or <code>.pdf."
        ),
    )
    parser.add_argument(
        "--code-mapping",
        type=Path,
        default=None,
        help="File mapping EVN session codes to legacy codes (e.g. E26A001 -> EC107).",
    )
    parser.add_argument(
        "--proposals",
        nargs="+",
        metavar="CODE",
        default=None,
        help="Send only for these proposal codes (useful for testing or resending).",
    )
    parser.add_argument(
        "--skip-compile",
        action="store_true",
        help=(
            "Skip LaTeX compilation even if .tex files in --pdf-dir have no "
            "corresponding PDF or are newer than their PDF."
        ),
    )
    parser.add_argument(
        "--sent-log",
        type=Path,
        default=None,
        metavar="FILE",
        help=(
            "JSON file used to track which emails have been sent.  Proposals "
            "already recorded in this log are skipped on re-runs.  Not written "
            "in --dry-run mode.  Defaults to <feedback-docx-stem>_sent.json."
        ),
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Send emails even for proposals already recorded in the sent log.",
    )
    parser.add_argument(
        "--subject-template",
        default=DEFAULT_SUBJECT,
        help=(
            "Email subject template.  Available tokens: {code}, {legacy_code}, {code_label}, {title}, {pi}. "
            f"Default: '{DEFAULT_SUBJECT}'."
        ),
    )
    parser.add_argument(
        "--from-address",
        help="Email address used in the From header.",
    )
    parser.add_argument(
        "--reply-to",
        help="Optional Reply-To address.",
    )
    parser.add_argument(
        "--cc",
        action="append",
        default=[],
        help="Additional CC recipients (may be supplied multiple times).",
    )
    parser.add_argument(
        "--smtp-server",
        default="smtp.gmail.com",
        help="SMTP server hostname (default: smtp.gmail.com).",
    )
    parser.add_argument(
        "--smtp-port",
        type=int,
        default=587,
        help="SMTP server port (default: 587).",
    )
    parser.add_argument(
        "--smtp-username",
        help="SMTP/IMAP username for authentication.",
    )
    parser.add_argument(
        "--smtp-password",
        help="SMTP/IMAP password (app password) for authentication.",
    )
    parser.add_argument(
        "--smtp-use-ssl",
        action="store_true",
        help="Use SMTP_SSL instead of STARTTLS.",
    )
    parser.add_argument(
        "--imap-server",
        default="imap.gmail.com",
        help="IMAP server used when saving drafts (default: imap.gmail.com).",
    )
    parser.add_argument(
        "--imap-port",
        type=int,
        default=993,
        help="IMAP SSL port (default: 993).",
    )
    parser.add_argument(
        "--drafts-folder",
        default="[Gmail]/Drafts",
        help="IMAP mailbox to append drafts to (default: [Gmail]/Drafts).",
    )

    # Delivery mode: dry-run (default) | --draft | --send
    mode_group = parser.add_mutually_exclusive_group()
    mode_group.add_argument(
        "--dry-run",
        action="store_true",
        default=True,
        help="Preview emails without sending or saving (default behaviour).",
    )
    mode_group.add_argument(
        "--draft",
        action="store_true",
        default=False,
        help=(
            "Save each email as a Gmail draft via IMAP instead of sending. "
            "Uses --smtp-username / --smtp-password and --imap-server."
        ),
    )
    mode_group.add_argument(
        "--send",
        action="store_true",
        default=False,
        help="Actually deliver the emails via SMTP.",
    )
    parser.add_argument(
        "--export-emails",
        type=Path,
        default=None,
        metavar="DIR",
        help="Save each email as a .eml file in this directory instead of (or in addition to) sending.",
    )
    return parser.parse_args(argv)


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def load_pi_emails(path: Path) -> Dict[str, str]:
    """Return {CODE_upper: email} from 'CODE: email' lines."""
    if not path.is_file():
        raise FileNotFoundError(f"PI emails file not found: {path}")
    result: Dict[str, str] = {}
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or ":" not in line:
            continue
        code, _, email = line.partition(":")
        email = email.strip()
        if email:
            result[code.strip().upper()] = email
    return result


def load_code_mapping(path: Optional[Path]) -> Dict[str, str]:
    """Return {EVN_code_upper: legacy_code}."""
    if path is None or not path.is_file():
        return {}
    text = path.read_text(encoding="utf-8", errors="replace")
    pattern = re.compile(
        r"\b([A-Z]\d{2}[A-Z]\d{3})\s+([A-Z]{1,2}\d{3,5})\b", re.IGNORECASE
    )
    mapping: Dict[str, str] = {}
    for m in pattern.finditer(text):
        evn = m.group(1).upper()
        if evn not in mapping:
            mapping[evn] = m.group(2)
    return mapping


# ---------------------------------------------------------------------------
# Docx parser
# ---------------------------------------------------------------------------

def _paragraph_to_html(para) -> str:
    """Convert a python-docx paragraph to an HTML fragment preserving bold/italic."""
    parts: List[str] = []
    for run in para.runs:
        text = _html.escape(run.text)
        if not text:
            continue
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    # Fall back to plain-escaped text if there are no runs (e.g. field codes).
    return "".join(parts) or _html.escape(para.text)


def parse_feedback_docx(path: Path) -> List[ProposalEmail]:
    """Extract per-proposal email content from the completed feedback .docx.

    Each Heading 1 paragraph marks the start of a new proposal section.
    Everything between two headings becomes the plain-text email body.
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise ImportError(
            "python-docx is required.  Install with: pip install python-docx"
        )

    doc = Document(str(path))
    proposals: List[ProposalEmail] = []
    current: Optional[ProposalEmail] = None

    def _finalise() -> None:
        if current is not None:
            proposals.append(current)

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text  # plain text, no markup

        if "heading 1" in style_name:
            _finalise()
            m = PROPOSAL_RE.search(text)
            code = m.group(0).upper() if m else ""
            # PI name follows the em-dash separator produced by build_email_docx
            pi_name = ""
            for sep in ("\u2014", "\u2013", "-"):
                if sep in text:
                    pi_name = text.split(sep, 1)[1].strip()
                    break
            current = ProposalEmail(code=code, legacy_code="", pi_name=pi_name)
            continue

        if current is None:
            # Before the first Heading 1 (TOC page) — skip.
            continue

        # Extract proposal title from the "Title:  <title>" paragraph.
        if not current.title and text.startswith("Title:"):
            current.title = text.split(":", 1)[1].strip()
            # Still include the paragraph in the body.

        current.paragraphs.append(text)
        current.html_paragraphs.append(_paragraph_to_html(para))

        # Detect unresolved placeholders in this paragraph.
        for m in UNRESOLVED_CHOICE_RE.finditer(text):
            current.unresolved_placeholders.append(m.group(0))
        if FEEDBACK_PLACEHOLDER in text:
            current.unresolved_placeholders.append(FEEDBACK_PLACEHOLDER)

    _finalise()
    return proposals


# ---------------------------------------------------------------------------
# Sent-log helpers
# ---------------------------------------------------------------------------

def load_sent_log(path: Path) -> Dict[str, Any]:
    """Return the sent log dict, or {} if the file doesn't exist yet."""
    if not path.is_file():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        print(f"[WARN] Could not read sent log {path}: {exc}", flush=True)
        return {}


def record_sent(
    log: Dict[str, Any],
    path: Path,
    code: str,
    email: str,
    status: str = "sent",
) -> None:
    """Add *code* to the sent log and persist it atomically.

    *status* is ``'sent'`` for SMTP delivery or ``'draft'`` for Gmail drafts.
    Only ``'sent'`` entries are skipped on re-runs; ``'draft'`` entries are
    informational only.
    """
    log[code] = {
        "email": email,
        "status": status,
        "sent_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
    }
    tmp = path.with_suffix(".tmp")
    try:
        tmp.write_text(json.dumps(log, indent=2), encoding="utf-8")
        tmp.replace(path)
    except OSError as exc:
        print(f"[WARN] Could not write sent log {path}: {exc}", flush=True)


# ---------------------------------------------------------------------------
# LaTeX compilation
# ---------------------------------------------------------------------------

def compile_latex_files(pdf_dir: Path) -> None:
    """Compile every .tex file in *pdf_dir* whose PDF is missing or stale.

    Runs ``pdflatex`` twice per file so cross-references resolve correctly.
    Compilation output is suppressed unless pdflatex exits with a non-zero
    status, in which case the last 20 lines of the log are printed.
    The working directory for each run is *pdf_dir* itself so that auxiliary
    files (.aux, .log) land alongside the sources.
    """
    pdflatex = shutil.which("pdflatex")
    if pdflatex is None:
        print(
            "[WARN] pdflatex not found on PATH — skipping LaTeX compilation. "
            "Install TeX Live or MiKTeX, or pass --skip-compile if PDFs already exist.",
            flush=True,
        )
        return

    tex_files = sorted(pdf_dir.glob("*.tex"))
    if not tex_files:
        return

    needs_compile: List[Path] = []
    for tex in tex_files:
        pdf = tex.with_suffix(".pdf")
        if not pdf.exists() or pdf.stat().st_mtime < tex.stat().st_mtime:
            needs_compile.append(tex)

    if not needs_compile:
        print(
            f"[INFO] All {len(tex_files)} PDF(s) in {pdf_dir} are up to date.",
            flush=True,
        )
        return

    print(
        f"[INFO] Compiling {len(needs_compile)} LaTeX file(s) in {pdf_dir} ...",
        flush=True,
    )

    cmd_base = [
        pdflatex,
        "-interaction=nonstopmode",
        "-halt-on-error",
        f"-output-directory={pdf_dir.resolve()}",
    ]

    for tex in needs_compile:
        label = tex.name
        ok = True
        for _pass in range(2):  # two passes for correct cross-references
            result = subprocess.run(
                cmd_base + [str(tex.resolve())],
                cwd=str(pdf_dir.resolve()),
                capture_output=True,
                text=True,
            )
            if result.returncode != 0:
                ok = False
                # Print the tail of the log to help diagnose the failure.
                log_lines = (result.stdout + result.stderr).splitlines()
                tail = "\n".join(log_lines[-20:])
                print(
                    f"[ERROR] pdflatex failed for {label} (pass {_pass + 1}):\n{tail}",
                    flush=True,
                )
                break

        if ok:
            print(f"  Compiled {label}", flush=True)


# ---------------------------------------------------------------------------
# PDF finder
# ---------------------------------------------------------------------------

def find_pdf(code: str, legacy_code: str, pdf_dir: Optional[Path]) -> Optional[Path]:
    """Return the path to the compiled feedback PDF, or None if not found."""
    if pdf_dir is None or not pdf_dir.is_dir():
        return None
    candidates = []
    if legacy_code:
        candidates.append(pdf_dir / f"{legacy_code}_{code}.pdf")
    candidates.append(pdf_dir / f"{code}.pdf")
    # Also search case-insensitively for robustness
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    # Fallback: glob for any pdf whose stem contains the code
    for pdf in sorted(pdf_dir.glob("*.pdf")):
        if code.upper() in pdf.stem.upper():
            return pdf
    return None


# ---------------------------------------------------------------------------
# Email delivery  (mirrors review_reminder.py)
# ---------------------------------------------------------------------------

def deliver_email(args: argparse.Namespace, message: EmailMessage) -> None:
    if args.smtp_use_ssl:
        server: smtplib.SMTP = smtplib.SMTP_SSL(args.smtp_server, args.smtp_port)
    else:
        server = smtplib.SMTP(args.smtp_server, args.smtp_port)
        server.ehlo()
        try:
            server.starttls()
            server.ehlo()
        except smtplib.SMTPException:
            pass
    try:
        if args.smtp_username and args.smtp_password:
            server.login(args.smtp_username, args.smtp_password)
        server.send_message(message)
    finally:
        server.quit()


def save_as_gmail_draft(args: argparse.Namespace, message: EmailMessage) -> None:
    """Append *message* to the Gmail Drafts folder via IMAP SSL.

    Uses the same username/password as the SMTP path so no extra credentials
    are needed.  The IMAP server and drafts folder are configurable via
    --imap-server and --drafts-folder.
    """
    if not args.smtp_username or not args.smtp_password:
        raise ValueError(
            "--smtp-username and --smtp-password are required to save drafts."
        )
    msg_bytes = message.as_bytes()
    with imaplib.IMAP4_SSL(args.imap_server, args.imap_port) as imap:
        imap.login(args.smtp_username, args.smtp_password)
        imap.append(
            args.drafts_folder,
            "\\Draft",
            imaplib.Time2Internaldate(time.time()),
            msg_bytes,
        )


def export_eml(
    message: EmailMessage,
    proposal: ProposalEmail,
    export_dir: Path,
) -> None:
    export_dir.mkdir(parents=True, exist_ok=True)
    slug = f"{proposal.legacy_code}_{proposal.code}" if proposal.legacy_code else proposal.code
    candidate = export_dir / f"{slug}.eml"
    idx = 1
    while candidate.exists():
        idx += 1
        candidate = export_dir / f"{slug}-{idx:02d}.eml"
    candidate.write_bytes(bytes(message))
    print(f"  Saved draft to {candidate}")


def build_message(
    proposal: ProposalEmail,
    to_email: str,
    pdf_path: Optional[Path],
    subject: str,
    args: argparse.Namespace,
) -> EmailMessage:
    sender = args.from_address or args.smtp_username or ""
    msg = EmailMessage()
    msg["Subject"] = subject
    if sender:
        msg["From"] = sender
    msg["To"] = to_email
    if args.reply_to:
        msg["Reply-To"] = args.reply_to
    if args.cc:
        msg["Cc"] = ", ".join(args.cc)

    # Plain-text fallback + HTML alternative (preserves bold/italic from docx).
    # EmailMessage.add_alternative() promotes the message to multipart/alternative;
    # a subsequent add_attachment() wraps that inside multipart/mixed automatically.
    msg.set_content(proposal.body_text)
    msg.add_alternative(proposal.html_body, subtype="html")

    if pdf_path is not None:
        pdf_bytes = pdf_path.read_bytes()
        msg.add_attachment(
            pdf_bytes,
            maintype="application",
            subtype="pdf",
            filename=pdf_path.name,
        )

    return msg


def send_proposal_email(
    proposal: ProposalEmail,
    to_email: str,
    pdf_path: Optional[Path],
    args: argparse.Namespace,
) -> None:
    code_label = f"{proposal.legacy_code}/{proposal.code}" if proposal.legacy_code else proposal.code
    subject = args.subject_template.format(
        code=proposal.code,
        legacy_code=proposal.legacy_code,
        code_label=code_label,
        title=proposal.title,
        pi=proposal.pi_name,
    )

    pdf_display = str(pdf_path) if pdf_path else "none"
    cc_display = f"  CC={', '.join(args.cc)}" if args.cc else ""

    # ---- Dry-run ----
    if not args.send and not args.draft:
        print(
            f"[DRY-RUN] {code_label} → {to_email}{cc_display}\n"
            f"  Subject: {subject}\n"
            f"  PDF: {pdf_display}\n"
            f"  Body preview ({len(proposal.body_text)} chars):\n"
            + "\n".join(f"    {line}" for line in proposal.body_text.splitlines()[:6])
            + ("\n    ..." if proposal.body_text.count("\n") > 6 else ""),
            flush=True,
        )
        if args.export_emails:
            msg = build_message(proposal, to_email, pdf_path, subject, args)
            export_eml(msg, proposal, args.export_emails)
        return

    sender = args.from_address or args.smtp_username
    if not sender:
        raise ValueError(
            "--from-address (or --smtp-username) is required."
        )

    msg = build_message(proposal, to_email, pdf_path, subject, args)

    # ---- Save as Gmail draft ----
    if args.draft:
        print(
            f"[DRAFT]   {code_label} → {args.drafts_folder} "
            f"(To: {to_email}  PDF: {pdf_display})",
            flush=True,
        )
        save_as_gmail_draft(args, msg)
        print(f"[SAVED]   {code_label} draft in {args.drafts_folder}", flush=True)
        if args.export_emails:
            export_eml(msg, proposal, args.export_emails)
        return

    # ---- Send via SMTP ----
    server_display = f"{args.smtp_server}:{args.smtp_port}"
    print(
        f"[SENDING] via {server_display}  {code_label} → {to_email}{cc_display}",
        flush=True,
    )
    deliver_email(args, msg)
    print(f"[SENT]    {code_label} ({to_email})", flush=True)

    if args.export_emails:
        export_eml(msg, proposal, args.export_emails)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    # Load inputs ---------------------------------------------------------
    if not args.feedback_docx.is_file():
        print(f"Feedback docx not found: {args.feedback_docx}", file=sys.stderr)
        return 1

    try:
        pi_emails = load_pi_emails(args.pi_emails_file)
    except (FileNotFoundError, ValueError) as exc:
        print(exc, file=sys.stderr)
        return 1

    code_mapping = load_code_mapping(args.code_mapping)

    try:
        proposals = parse_feedback_docx(args.feedback_docx)
    except (ImportError, Exception) as exc:
        print(f"Error reading feedback docx: {exc}", file=sys.stderr)
        return 1

    if not proposals:
        print("No proposals found in the feedback docx.", file=sys.stderr)
        return 1

    print(f"Parsed {len(proposals)} proposals from {args.feedback_docx.name}.")

    # Compile LaTeX → PDF if needed --------------------------------------
    if args.pdf_dir and args.pdf_dir.is_dir() and not args.skip_compile:
        compile_latex_files(args.pdf_dir)

    # Attach legacy codes from mapping ------------------------------------
    for p in proposals:
        p.legacy_code = code_mapping.get(p.code, "")

    # Sent log ------------------------------------------------------------
    sent_log_path = args.sent_log or args.feedback_docx.with_name(
        args.feedback_docx.stem + "_sent.json"
    )
    sent_log = load_sent_log(sent_log_path)
    if sent_log:
        already = sorted(sent_log)
        print(
            f"[LOG] {len(already)} proposal(s) already sent per {sent_log_path.name}: "
            + ", ".join(already)
        )

    # Filter to requested proposals if specified --------------------------
    filter_codes: Optional[Set[str]] = None
    if args.proposals:
        filter_codes = {c.upper() for c in args.proposals}

    # Process each proposal -----------------------------------------------
    skipped_placeholders = 0
    skipped_no_email = 0
    skipped_already_sent = 0
    sent = 0
    errors = 0

    for proposal in proposals:
        if filter_codes and proposal.code not in filter_codes:
            continue

        code_label = (
            f"{proposal.legacy_code}/{proposal.code}"
            if proposal.legacy_code else proposal.code
        )

        # Skip proposals already successfully sent (unless --force).
        # Draft mode does not count as "sent" — only --send does.
        if not args.force and proposal.code in sent_log:
            entry = sent_log[proposal.code]
            if entry.get("status") == "sent":
                print(
                    f"[SKIP] {code_label}: already sent to {entry['email']} "
                    f"at {entry['sent_at']} — use --force to resend.",
                    flush=True,
                )
                skipped_already_sent += 1
                continue

        # Skip proposals with unresolved placeholders.
        if proposal.unresolved_placeholders:
            unique = list(dict.fromkeys(proposal.unresolved_placeholders))
            print(
                f"[SKIP] {code_label}: {len(unique)} unresolved placeholder(s) — "
                + ", ".join(f'"{p}"' for p in unique[:3])
                + ("" if len(unique) <= 3 else f" (+{len(unique)-3} more)"),
                flush=True,
            )
            skipped_placeholders += 1
            continue

        to_email = pi_emails.get(proposal.code)
        if not to_email:
            print(
                f"[SKIP] {code_label}: no PI email found — add it to {args.pi_emails_file.name}",
                flush=True,
            )
            skipped_no_email += 1
            continue

        pdf_path = find_pdf(proposal.code, proposal.legacy_code, args.pdf_dir)
        if args.pdf_dir and pdf_path is None:
            print(
                f"[WARN] {code_label}: no PDF found in {args.pdf_dir} — sending without attachment.",
                flush=True,
            )

        try:
            send_proposal_email(proposal, to_email, pdf_path, args)
            sent += 1
            # Record in the log only for real actions (not dry-run).
            if args.send:
                record_sent(sent_log, sent_log_path, proposal.code, to_email, status="sent")
            elif args.draft:
                record_sent(sent_log, sent_log_path, proposal.code, to_email, status="draft")
        except ValueError as exc:
            print(f"[ERROR] {code_label}: {exc}", file=sys.stderr)
            errors += 1
        except (smtplib.SMTPException, imaplib.IMAP4.error) as exc:
            print(f"[ERROR] {code_label}: {exc}", file=sys.stderr)
            errors += 1

    # Summary -------------------------------------------------------------
    print()
    if args.send:
        mode = "sent"
    elif args.draft:
        mode = "saved as draft"
    else:
        mode = "dry-run preview"
    print(
        f"Done. {mode}: {sent}  |  skipped (already sent): {skipped_already_sent}  "
        f"|  skipped (placeholders): {skipped_placeholders}  "
        f"|  skipped (no email): {skipped_no_email}  |  errors: {errors}"
    )
    if (args.send or args.draft) and sent_log:
        print(f"Sent log: {sent_log_path}")
    return 0 if errors == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
